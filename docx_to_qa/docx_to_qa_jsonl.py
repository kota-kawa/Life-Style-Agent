#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
docx_to_qa_jsonl.py

OpenAI の Python SDK 経由で Google Gemini（OpenAI 互換エンドポイント）を呼び出し、
.docx から本文を抽出 → チャンク分割 → 各チャンクから網羅的な Q/A を生成 →
question, answer の2キーのみを持つ JSONL に出力するスクリプト。

■ 主な特徴
- secrets.env から API キー等を読み込み（python-dotenv）
- python-docx で段落・表セルのテキストを抽出
- 安全なチャンク分割（オーバーラップあり）
- モデルには「外部知識を使わず、原文のみから Q/A を網羅的に作る」プロンプトを付与
- 応答の JSON を堅牢にパース（フェンス/余分な文除去、再試行リカバリ）
- 重複 Q を避けるための正規化 & 既出検知
- 再開可能な簡易ステート（--state で保存/読み込み）
- 並列実行（--workers 指定）

■ 使い方（例）
1) 必要パッケージのインストール:
   pip install openai python-dotenv python-docx tenacity tqdm

2) secrets.env を作成（最低限）:
   GEMINI_API_KEY=あなたのGoogle AI StudioのAPIキー
   OPENAI_BASE_URL=https://generativelanguage.googleapis.com/v1beta/openai/
   GEMINI_MODEL=gemini-2.5-pro

3) 実行:
   python docx_to_qa_jsonl.py input.docx output.jsonl

   オプション:
   --chunk-chars 8000        # チャンクの文字数上限（既定: 8000）
   --chunk-overlap 400       # チャンクの重なり（既定: 400）
   --workers 3               # 併行ワーカー数（既定: 2）
   --state qa_state.json     # 途中経過の保存/再開用ステートファイル
   --max-retries 3           # APIリトライ回数（既定: 3）
   --temperature 0.2         # 生成温度（既定: 0.2）
   --model gemini-2.5-pro    # secrets.env の GEMINI_MODEL より優先
   --rate-wait 0.0           # 呼び出し間の待機秒（既定: 0）

■ 注意
- 出力 JSONL は各行 {"question": "...", "answer": "..."} のみ。
- 画像や数式埋め込みは python-docx の仕様上テキスト化できない場合があります。
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
import hashlib
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Iterable, Tuple, Set

from dotenv import load_dotenv
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from tqdm import tqdm

# docx
from docx import Document  # python-docx

_ENV_FILES = [
    Path(__file__).resolve().parent / "secrets.env",
    Path(__file__).resolve().parents[1] / "secrets.env",
]
for env_path in _ENV_FILES:
    load_dotenv(env_path, override=False)

# OpenAI SDK（OpenAI互換エンドポイントでGeminiを叩く）
from openai import OpenAI
from openai import APIError, RateLimitError, APITimeoutError, InternalServerError


# =========================
# 入出力・設定
# =========================

@dataclass
class Settings:
    api_key: str
    base_url: str
    model: str
    temperature: float = 0.2
    max_retries: int = 3
    rate_wait: float = 0.0  # 連続呼び出し間のスリープ秒（雑なレート制御）


SYSTEM_PROMPT = (
    "You are a meticulous annotator that converts source text into exhaustive Q/A pairs.\n"
    "Constraints:\n"
    "1) Use ONLY the provided text; do NOT add external knowledge.\n"
    "2) Cover ALL atomic facts, definitions, enumerations, properties, constraints, procedures, and relationships present in the text.\n"
    "3) Keep each question focused on ONE fact/topic whenever possible.\n"
    "4) Answers must be self-contained, precise, and directly supported by the text.\n"
    "5) Output STRICT JSON with this schema ONLY: [{\"question\": str, \"answer\": str}, ...]. No preface, no trailing notes.\n"
    "6) Preserve the language of the source text (e.g., Japanese text → Japanese Q/A).\n"
)

USER_PROMPT_TEMPLATE = (
    "Convert the following SOURCE TEXT into an exhaustive set of Q/A pairs.\n"
    "Return ONLY a JSON array of objects with keys 'question' and 'answer'.\n\n"
    "SOURCE TEXT:\n"
    "{chunk}\n"
)


# =========================
# .docx 読み込み
# =========================

def extract_text_from_docx(path: str) -> str:
    """
    .docx からテキストを抽出する。
    段落・表セル内容を順番に結合。空行は圧縮。
    """
    doc = Document(path)
    parts: List[str] = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 表（セルごと）
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    # 空行連続を抑制しつつ結合
    text = "\n".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# =========================
# チャンク分割
# =========================

def chunk_text(text: str, max_chars: int = 8000, overlap: int = 400) -> List[str]:
    """
    文字数ベースで素朴にチャンク分割（UTF-8の見かけ文字数）。
    overlap で文脈を少し重ねる。
    """
    if max_chars <= 0:
        raise ValueError("max_chars must be > 0")
    if overlap < 0 or overlap >= max_chars:
        raise ValueError("overlap must be >= 0 and < max_chars")

    n = len(text)
    chunks = []
    start = 0
    while start < n:
        end = min(start + max_chars, n)
        chunk = text[start:end]
        chunks.append(chunk.strip())
        if end >= n:
            break
        start = end - overlap  # 次の開始位置（オーバーラップ分戻す）
    return chunks


# =========================
# OpenAI SDK 経由で Gemini を呼ぶ
# =========================

class GeminiClient:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.client = OpenAI(api_key=settings.api_key, base_url=settings.base_url)

    @retry(
        reraise=True,
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=10),
        retry=retry_if_exception_type((APIError, RateLimitError, APITimeoutError, InternalServerError))
    )
    def qa_from_chunk(self, chunk: str) -> List[Dict[str, str]]:
        """
        1チャンクから Q/A 配列を生成して返す。
        レスポンスは厳密な JSON 配列を期待し、乱れた場合は整形を試みる。
        """
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": USER_PROMPT_TEMPLATE.format(chunk=chunk)}
        ]

        resp = self.client.chat.completions.create(
            model=self.settings.model,
            temperature=self.settings.temperature,
            messages=messages,
        )

        if self.settings.rate_wait > 0:
            time.sleep(self.settings.rate_wait)

        content = resp.choices[0].message.content if resp and resp.choices else ""
        return self._parse_json_array(content, fallback_repair=True)

    @staticmethod
    def _parse_json_array(text: str, fallback_repair: bool = True) -> List[Dict[str, str]]:
        """
        応答テキストから JSON 配列を抽出・ロードする。
        - コードフェンスや前置き/後置きが混入した場合もできる限り復元。
        - 各要素は {"question": str, "answer": str} のみを許容。
        """
        if not text:
            raise ValueError("Empty response from model.")

        # ```json ... ``` などのフェンス除去
        fence = re.search(r"```(?:json)?\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
        candidate = fence.group(1) if fence else text

        # 先頭の [ と最後の ] のペアを一番外側で検出
        start_idx = candidate.find("[")
        end_idx = candidate.rfind("]")
        if start_idx == -1 or end_idx == -1 or end_idx <= start_idx:
            if not fallback_repair:
                raise ValueError("Failed to locate JSON array in response.")
            # ゆるい修復：オブジェクト行を拾って配列化を試みる
            objs = re.findall(r"\{.*?\}", candidate, flags=re.DOTALL)
            if not objs:
                raise ValueError("No JSON objects found in response.")
            try:
                items = [json.loads(o) for o in objs]
            except json.JSONDecodeError as e:
                raise ValueError(f"JSON decode error during repair: {e}") from e
            cleaned = [GeminiClient._clean_item(it) for it in items if isinstance(it, dict)]
            return [it for it in cleaned if "question" in it and "answer" in it]

        json_str = candidate[start_idx:end_idx + 1]

        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            if not fallback_repair:
                raise
            # 末尾カンマなど軽微な崩れを修復
            json_str2 = re.sub(r",\s*]", "]", json_str)
            data = json.loads(json_str2)

        if not isinstance(data, list):
            raise ValueError("Parsed data is not a list.")

        cleaned = []
        for it in data:
            if isinstance(it, dict):
                obj = GeminiClient._clean_item(it)
                if "question" in obj and "answer" in obj:
                    cleaned.append(obj)
        return cleaned

    @staticmethod
    def _clean_item(it: Dict[str, Any]) -> Dict[str, str]:
        """
        キー厳格化：「question」「answer」以外は捨てる。値は文字列化。
        """
        q = it.get("question", "")
        a = it.get("answer", "")
        return {
            "question": str(q).strip(),
            "answer": str(a).strip(),
        }


# =========================
# 重複検知・正規化
# =========================

def normalize_question(q: str) -> str:
    """
    質問文の正規化：小文字化・空白圧縮・句読点の一部正規化など。
    日本語もあるため強すぎる正規化は避ける。
    """
    s = q.strip()
    s = re.sub(r"\s+", " ", s)
    return s.lower()


# =========================
# ステート（再開用）
# =========================

@dataclass
class RunState:
    processed: Set[int]  # チャンク index の集合

    def to_json(self) -> Dict[str, Any]:
        return {"processed": sorted(list(self.processed))}

    @staticmethod
    def from_json(d: Dict[str, Any]) -> "RunState":
        proc = set(d.get("processed", []))
        return RunState(processed=proc)


def load_state(path: str | None) -> RunState:
    if not path or not os.path.exists(path):
        return RunState(processed=set())
    with open(path, "r", encoding="utf-8") as f:
        return RunState.from_json(json.load(f))


def save_state(path: str | None, state: RunState) -> None:
    if not path:
        return
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(state.to_json(), f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


# =========================
# メイン処理
# =========================

def build_settings(args: argparse.Namespace) -> Settings:
    for env_path in _ENV_FILES:
        load_dotenv(env_path, override=False)

    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("OPENAI_API_KEY") or ""
    if args.api_key:
        api_key = args.api_key
    if not api_key:
        print(
            "ERROR: API key not found. Set GEMINI_API_KEY (or OPENAI_API_KEY) in secrets.env or pass --api-key.",
            file=sys.stderr,
        )
        sys.exit(1)

    base_url = os.getenv("OPENAI_BASE_URL", "https://generativelanguage.googleapis.com/v1beta/openai/")
    if args.base_url:
        base_url = args.base_url

    model = args.model or os.getenv("GEMINI_MODEL", "gemini-2.5-pro")

    return Settings(
        api_key=api_key,
        base_url=base_url,
        model=model,
        temperature=args.temperature,
        max_retries=args.max_retries,
        rate_wait=args.rate_wait,
    )


def write_jsonl(path: str, items: Iterable[Dict[str, str]]) -> int:
    count = 0
    with open(path, "a", encoding="utf-8") as f:
        for it in items:
            # question, answer のみで書く（その他フィールドは持たせない）
            rec = {"question": it["question"], "answer": it["answer"]}
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
            count += 1
    return count


def process_docx_to_jsonl(
    in_docx: str,
    out_jsonl: str,
    settings: Settings,
    chunk_chars: int = 8000,
    chunk_overlap: int = 400,
    workers: int = 2,
    state_path: str | None = None,
) -> None:
    """
    .docx → JSONL 変換のトップレベル処理。
    """
    text = extract_text_from_docx(in_docx)
    if not text:
        print("No text extracted from DOCX.", file=sys.stderr)
        return

    chunks = chunk_text(text, max_chars=chunk_chars, overlap=chunk_overlap)
    total_chunks = len(chunks)
    print(f"Extracted {total_chunks} chunk(s).")

    # 既存の出力がある場合も追記安全にするため、重複 Q 検知セットを用意
    seen_questions: Set[str] = set()

    # ステートロード
    state = load_state(state_path)

    client = GeminiClient(settings)

    # 並列実行
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def task(idx: int, chunk: str) -> Tuple[int, List[Dict[str, str]]]:
        qa = client.qa_from_chunk(chunk)
        return idx, qa

    pending_indices = [i for i in range(total_chunks) if i not in state.processed]

    with ThreadPoolExecutor(max_workers=max(1, workers)) as ex, tqdm(total=len(pending_indices), desc="Chunks") as pbar:
        future_map = {ex.submit(task, i, chunks[i]): i for i in pending_indices}

        with open(out_jsonl, "a", encoding="utf-8") as fout:
            for future in as_completed(future_map):
                idx = future_map[future]
                try:
                    _, qa_list = future.result()
                except Exception as e:
                    print(f"[Chunk {idx}] ERROR: {e}", file=sys.stderr)
                    # 失敗したチャンクは未処理のまま残す
                    pbar.update(1)
                    continue

                # 重複を避けつつ書き込み
                written = 0
                for item in qa_list:
                    q_norm = normalize_question(item["question"])
                    if not q_norm or q_norm in seen_questions:
                        continue
                    seen_questions.add(q_norm)
                    rec = {"question": item["question"], "answer": item["answer"]}
                    fout.write(json.dumps(rec, ensure_ascii=False) + "\n")
                    written += 1

                state.processed.add(idx)
                save_state(state_path, state)
                pbar.update(1)
                tqdm.write(f"[Chunk {idx}] wrote {written} Q/A")

    print("Done.")


# =========================
# CLI
# =========================

def main():
    parser = argparse.ArgumentParser(description="Convert DOCX to exhaustive Q/A JSONL via Gemini (OpenAI-compatible).")
    parser.add_argument("input_docx", help="入力 .docx パス")
    parser.add_argument("output_jsonl", help="出力 .jsonl パス（追記）")
    parser.add_argument("--chunk-chars", type=int, default=8000, help="チャンク文字数上限（既定: 8000）")
    parser.add_argument("--chunk-overlap", type=int, default=400, help="チャンク重なり文字数（既定: 400）")
    parser.add_argument("--workers", type=int, default=2, help="並列ワーカー数（既定: 2）")
    parser.add_argument("--state", dest="state_path", default=None, help="再開用ステートファイルパス（例: qa_state.json）")
    parser.add_argument("--temperature", type=float, default=0.2, help="生成温度（既定: 0.2）")
    parser.add_argument("--max-retries", type=int, default=3, help="API リトライ回数（既定: 3）")
    parser.add_argument("--rate-wait", type=float, default=0.0, help="各API呼び出し後の待機秒（既定: 0）")
    parser.add_argument("--api-key", default=None, help="API キー（secrets.env より優先）")
    parser.add_argument("--base-url", default=None, help="OpenAI 互換エンドポイント（secrets.env より優先）")
    parser.add_argument("--model", default=None, help="モデル名（secrets.env の GEMINI_MODEL より優先）")

    args = parser.parse_args()
    settings = build_settings(args)

    process_docx_to_jsonl(
        in_docx=args.input_docx,
        out_jsonl=args.output_jsonl,
        settings=settings,
        chunk_chars=args.chunk_chars,
        chunk_overlap=args.chunk_overlap,
        workers=args.workers,
        state_path=args.state_path,
    )


if __name__ == "__main__":
    main()
